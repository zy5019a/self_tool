from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt
from datetime import datetime, timedelta
from langchain_community.chat_models import ChatOpenAI
from langchain_core.prompts import PromptTemplate
from langchain_experimental.sql import SQLDatabaseSequentialChain
from langchain.sql_database import SQLDatabase
from pydantic import BaseModel, Field
import json
import re
from minio import Minio
import io
import requests

from server.agent.tools.sql_search import UserQueryForSQL


class TemplateInput(BaseModel):
    doc_path: str = Field(description="The download url to the Word document to be processed.")


# MinIO 客户端初始化
minio_client = Minio(
    "127.0.0.1:9000",
    access_key="minioadmin",
    secret_key="minioadmin",
    secure=False
)

# 确保存储桶存在
bucket_name = "template"
if not minio_client.bucket_exists(bucket_name):
    minio_client.make_bucket(bucket_name)


class TableFiller:
    def __init__(self, minio_client, bucket_name):
        self.minio_client = minio_client
        self.bucket_name = bucket_name

    def clean_response(self, response):
        cleaned_response = response.strip()
        if not cleaned_response.startswith("{"):
            cleaned_response = "{" + cleaned_response
        if not cleaned_response.endswith("}"):
            cleaned_response = cleaned_response + "}"
        cleaned_response = re.sub(r'\bNone\b', '""', cleaned_response)
        return cleaned_response

    def is_merged_cell_row(self, row):
        for cell in row.cells:
            grid_span = cell._element.xpath('.//w:gridSpan')
            if grid_span and int(
                    grid_span[0].get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')) > 1:
                return True
        return False

    def get_table_name(self, doc, table):
        for block in doc.element.body:
            if block.tag.endswith('tbl'):
                if block == table._element:
                    prev = block.getprevious()
                    while prev is not None and not prev.tag.endswith('p'):
                        prev = prev.getprevious()
                    if prev is not None and prev.tag.endswith('p'):
                        return prev.text
        return "Unknown Table Name"

    def run_sql_query(self, natural_language_query, db_user='root', db_password='root', db_host='localhost',
                      db_name='chart'):
        connection_str = f"mysql+pymysql://{db_user}:{db_password}@{db_host}/{db_name}"
        db = SQLDatabase.from_uri(connection_str)
        llm = ChatOpenAI(openai_api_key="", model="",
                         base_url="")
        query_prompt_prefix = """Only use the following tables:
        {table_info}
        Question: {input}"""
        query_prompt_mysql = """You are a MySQL expert. Given an input question, first create a syntactically correct MySQL query to run, then look at the results of the query and return the answer to the input question.
        Unless the user specifies in the question a specific number of examples to obtain, query for at most {top_k} results using the LIMIT clause as per MySQL. You can order the results to return the most informative data in the database.
        Never query for all columns from a table. You must query only the columns that are needed to answer the question. Wrap each column name in backticks (`) to denote them as delimited identifiers.
        Pay attention to use only the column names you can see in the tables below. Be careful to not query for columns that do not exist. Also, pay attention to which column is in which table.
        Pay attention to use CURDATE() function to get the current date, if the question involves "today". If you need to perform multiple queries, please use the JOIN method to combine them into a single query.
        Please make sure you proofread your generated statements, no syntax error, no semantic error, no logic error, no query that doesn't fit the question, no result that doesn't fit the question, no answer that doesn't fit the question.
        And your SQL Query must be output directly as a string, there can't be any characters outside of the query in your query, all the query must be output directly in the query, don't add anything outside of the query, don't say any extra one character outside of the query.
        Additionally, the generated query must be a single continuous query without multiple semicolons.在你的query中不允许使用已知字段外的column名，任何情况都不允许order by排序。
        Use the following format:
        Question: Question here
        SQLQuery: SQL Query to run, only the query, no other characters
        SQLResult: Result of the SQLQuery
        Answer: Final answer here
        Please generate a pure text format SQL query without any code block markers."""
        query_prompt = PromptTemplate(
            input_variables=["input", "table_info", "top_k"],
            template=query_prompt_mysql + query_prompt_prefix,
        )
        db_chain = SQLDatabaseSequentialChain.from_llm(llm, db, query_prompt=query_prompt, verbose=True, top_k='10000')

        max_retries = 5
        for attempt in range(max_retries):
            try:
                response = db_chain.run(natural_language_query)
                cleaned_response = self.clean_response(response)
                response_dict = json.loads(cleaned_response)
                return response_dict
            except (json.JSONDecodeError, Exception) as e:
                print(f"Error occurred on attempt {attempt + 1}: {e}")
                if attempt == max_retries - 1:
                    raise e
                else:
                    continue

    def fill_and_extend_table(self, docx_stream, output_path=None, data=None):
        doc = Document(docx_stream)
        table_info = []
        for table in doc.tables:
            headers = []
            start_row = 0
            table_name = self.get_table_name(doc, table)
            if len(table.rows) > 1 and self.is_merged_cell_row(table.rows[0]):
                headers = [cell.text.strip() for cell in table.rows[1].cells]
                start_row = 2
            else:
                headers = [cell.text.strip() for cell in table.rows[0].cells]
                start_row = 1
            if data and table_name in data:
                table_data = data[table_name]
                if all(column in headers for column in table_data.keys() if column != "总结"):
                    first_empty_row = None
                    for row_index, row in enumerate(table.rows[start_row:], start=start_row):
                        if all(cell.text.strip() == '' for cell in row.cells):
                            first_empty_row = row_index
                            break
                    if first_empty_row is None:
                        first_empty_row = len(table.rows)
                    sequence_number = 1
                    for row_data in zip(*table_data.values()):
                        if first_empty_row < len(table.rows):
                            row = table.rows[first_empty_row]
                        else:
                            row = table.add_row()
                        for cell_index, column in enumerate(headers):
                            if column == "序号":
                                row.cells[cell_index].text = str(sequence_number)
                                sequence_number += 1
                            elif column in table_data:
                                row.cells[cell_index].text = str(row_data[headers.index(column) - 1])
                        first_empty_row += 1
                    if "总结" in table_data:
                        summary_row = table.add_row()
                        summary_cell = summary_row.cells[0]
                        summary_cell.merge(summary_row.cells[-1])
                        run = summary_cell.paragraphs[0].add_run(table_data["总结"])
                        run.font.name = '宋体'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                        run.font.size = Pt(12)
            else:
                columns_without_index = [header for header in headers if header != "序号"]
                table_info.append({
                    "Table Name": table_name,
                    "Columns": columns_without_index
                })
        if output_path and data:
            doc.save(output_path)
            xml_buffer = io.BytesIO()
            doc.save(xml_buffer)
            xml_buffer.seek(0)
            self.minio_client.put_object(self.bucket_name, output_path, xml_buffer, len(xml_buffer.getvalue()),
                                         content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
            download_link = self.minio_client.presigned_get_object(self.bucket_name, output_path,
                                                                   expires=timedelta(hours=1))
            return f"Download link: {download_link}"
        else:
            now = datetime.now()
            current_year = now.year
            current_month = now.month
            request_string = f"我需要下列一个或者多个表的 {current_year} 年 5 月全部数据：\n"
            for info in table_info:
                request_string += f"表名：{info['Table Name']}\n列名：{', '.join(info['Columns'])}\n\n"
            request_string += '''请在执行完语句后的回答中提供全部数据。
            以下是传回来的其中一张表的数据，你可以参考这个格式来回答：
            {
                "表XXX": {
                "XX号": ["A001", "A002", "A003"],
                "XX号": ["B001", "B002", "B003"],
                "XX时间": ["2024-01-01", "2024-02-01", "2024-03-01"],
                "XX时间": ["2024-01-10", "2024-02-10", "2024-03-10"],
                "XX": ["1", "2", "3"],
                "XX时间": ["2024-04-01", "2024-05-01", "2024-06-01"],
                "备注": ["无", "无", "无"],
                "总结": "这是总结内容，包含整体的描述和分析。"
                }
            }，这种是在Table Name只有一张表XXX的情况下的示例，并且每个key不允许用省略号代替任何数据，一定要写全。另外总结是每个表都有一个总结，如果有多张表，data内就需要多个这样的数据结构。
            请直接输出这种格式，不要在{}外面加任何东西例如“json{}”，不要说{}以外任何多余的一个字符
            '''
            print(request_string)
            return request_string


def process_document(doc_path: str):
    table_filler = TableFiller(minio_client, bucket_name)

    response = requests.get(doc_path)
    docx_stream = io.BytesIO(response.content)

    request_string = table_filler.fill_and_extend_table(docx_stream)
    data_query = UserQueryForSQL(natural_language_query=request_string)
    data = table_filler.run_sql_query(data_query.natural_language_query)
    output_path = 'filled_template.docx'
    return table_filler.fill_and_extend_table(docx_stream, output_path, data)

